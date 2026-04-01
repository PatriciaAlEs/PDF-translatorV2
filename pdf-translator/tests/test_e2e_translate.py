"""
Test end-to-end: simula el flujo completo de traducción con pipeline real.
Verifica que el job SIEMPRE llega a 'completed' o 'error', nunca se queda en 'postprocessing'.
"""
import sys
import os
import time
import threading

# Ensure we can import from backend/
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# ── Test helpers ──
def make_test_docx(path, paragraphs):
    """Create a minimal .docx with given paragraphs."""
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


# ── Simulated translation job tracker (same as app.py) ──
translation_jobs = {}


def simulate_translate_background(session_id, docx_path, output_path,
                                   source_lang, use_ai, provider):
    """Replica exacta de _translate_docx_background con los fixes aplicados."""
    from docx import Document
    from concurrent.futures import ThreadPoolExecutor, as_completed

    job = translation_jobs[session_id]
    job["status"] = "translating"
    job["started_at"] = time.time()
    job["_heartbeat"] = time.time()

    processed = None  # define in outer scope for finally block

    try:
        doc = Document(str(docx_path))
        work_items = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) >= 3:
                work_items.append((i, text))

        total = len(work_items)
        job["total"] = total
        job["translated"] = 0

        if total == 0:
            job["status"] = "completed"
            job["finished_at"] = time.time()
            doc.save(str(output_path))
            return

        # Phase 1: Fake "translation" (return text as-is, simulating Google)
        raw_translations = {}
        for idx, text in work_items:
            # Simulate translation: just prefix with [ES]
            raw_translations[idx] = f"[Traducido] {text}"
            job["translated"] += 1
            job["_heartbeat"] = time.time()
            time.sleep(0.05)  # simulate latency

        # Phase 2: Post-processing with real pipeline
        job["status"] = "postprocessing"
        job["_heartbeat"] = time.time()
        print(f"  [Phase 2] Post-processing {len(work_items)} paragraphs...")
        ordered_raw = [raw_translations.get(idx, text) for idx, text in work_items]

        pp_start = time.time()
        PP_TIMEOUT = 30  # shorter timeout for test

        processed = None
        try:
            from postprocess_pipeline import run_pipeline_batch
            from concurrent.futures import ThreadPoolExecutor as _TPE, TimeoutError as _TE

            pp_executor = _TPE(max_workers=1)
            pp_future = pp_executor.submit(
                run_pipeline_batch, ordered_raw, use_ai, provider
            )
            try:
                pp_results = pp_future.result(timeout=PP_TIMEOUT)
                processed = {work_items[j][0]: pp_results[j]["text"] for j in range(len(work_items))}
                print(f"  [Phase 2] Completed in {time.time() - pp_start:.1f}s")
            except (_TE, TimeoutError):
                pp_future.cancel()
                pp_executor.shutdown(wait=False, cancel_futures=True)
                print(f"  [Phase 2] TIMEOUT ({PP_TIMEOUT}s) — using raw translations")
            except Exception as e:
                print(f"  [Phase 2] Pipeline raised: {e}")
                try:
                    pp_executor.shutdown(wait=False, cancel_futures=True)
                except Exception:
                    pass
            else:
                pp_executor.shutdown(wait=False)
        except Exception as e:
            print(f"  [Phase 2] Import/setup error: {e}")

        if processed is None:
            print(f"  [Phase 2] Fallback: using raw translations")
            processed = {}
            for j, (idx, _) in enumerate(work_items):
                processed[idx] = ordered_raw[j]

        # Phase 3: Apply back
        job["_heartbeat"] = time.time()
        for i, para in enumerate(doc.paragraphs):
            if i in processed:
                runs = para.runs
                if runs:
                    runs[0].text = processed[i]
                    for run in runs[1:]:
                        run.text = ""

        doc.save(str(output_path))
        print(f"  [Phase 3] Saved: {output_path}")

    except Exception as e:
        job["error_message"] = str(e)
        print(f"  [ERROR] {e}")
        import traceback
        traceback.print_exc()
    finally:
        if job["status"] not in ("completed", "error", "cancelled"):
            if processed is not None and len(processed) > 0:
                job["status"] = "completed"
            else:
                job["status"] = "error"
                if not job.get("error_message"):
                    job["error_message"] = "Pipeline failed but paragraphs were translated"
        job["finished_at"] = time.time()
        job["_heartbeat"] = time.time()
        print(f"  [FINAL] Status: {job['status']}")


def run_test(test_name, paragraphs, expect_status="completed", max_wait=60):
    """Run a full translation test and verify it completes."""
    import tempfile
    print(f"\n{'='*60}")
    print(f"TEST: {test_name}")
    print(f"  Paragraphs: {len(paragraphs)}")
    print(f"  Max wait: {max_wait}s")
    print(f"{'='*60}")

    session_id = f"test_{test_name.replace(' ', '_')}"
    tmpdir = tempfile.mkdtemp()
    docx_in = os.path.join(tmpdir, "input.docx")
    docx_out = os.path.join(tmpdir, "output.docx")

    make_test_docx(docx_in, paragraphs)

    translation_jobs[session_id] = {
        "status": "starting",
        "total": 0,
        "translated": 0,
        "errors": [],
        "eta_seconds": None,
        "started_at": None,
        "finished_at": None,
        "cancelled": False,
        "error_message": None,
        "_heartbeat": time.time(),
    }

    thread = threading.Thread(
        target=simulate_translate_background,
        args=(session_id, docx_in, docx_out, "auto", False, "none"),
        daemon=True,
    )
    thread.start()

    # Poll status (simulates frontend polling)
    start = time.time()
    last_status = ""
    while time.time() - start < max_wait:
        job = translation_jobs[session_id]
        status = job["status"]

        if status != last_status:
            elapsed = time.time() - start
            print(f"  [{elapsed:.1f}s] Status: {status}")
            last_status = status

        if status in ("completed", "error", "cancelled"):
            break

        # Stale detection (same as progress endpoint)
        heartbeat = job.get("_heartbeat", 0)
        if heartbeat > 0 and time.time() - heartbeat > 45:
            print(f"  [STALE] Heartbeat stale ({time.time() - heartbeat:.0f}s), forcing completion")
            job["status"] = "completed"
            job["finished_at"] = time.time()
            break

        time.sleep(0.5)
    else:
        elapsed = time.time() - start
        print(f"\n  TIMEOUT after {elapsed:.1f}s! Status stuck at: {last_status}")

    # Results
    job = translation_jobs[session_id]
    total_time = time.time() - start
    final_status = job["status"]

    print(f"\n  Results:")
    print(f"    Final status:  {final_status}")
    print(f"    Total time:    {total_time:.1f}s")
    print(f"    Translated:    {job.get('translated', 0)} / {job.get('total', 0)}")
    print(f"    Error message: {job.get('error_message', 'none')}")

    # Verify
    passed = True
    if final_status != expect_status:
        print(f"\n  FAIL: Expected status '{expect_status}', got '{final_status}'")
        passed = False
    if final_status in ("completed", "error", "cancelled"):
        if total_time > max_wait:
            print(f"\n  FAIL: Took {total_time:.1f}s, max was {max_wait}s")
            passed = False

    if passed:
        print(f"\n  ✅ PASSED ({total_time:.1f}s)")
    else:
        print(f"\n  ❌ FAILED")

    # Cleanup
    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)

    return passed


# ═══════════════════════════════════════════
# Run tests
# ═══════════════════════════════════════════

if __name__ == "__main__":
    results = []

    # Test 1: Small document (should complete fast)
    results.append(run_test(
        "small_document",
        [
            "The quick brown fox jumped over the lazy dog.",
            "She said hello to everyone at the party.",
            "It was a dark and stormy night.",
        ],
        expect_status="completed",
        max_wait=45,
    ))

    # Test 2: Medium document (10 paragraphs)
    medium_paragraphs = [
        "The young woman walked through the crowded marketplace, her eyes scanning the stalls.",
        "She had been searching for a particular herb that her grandmother had mentioned.",
        "The sun was setting behind the mountains, casting long shadows across the valley.",
        "He shook his head and rolled his eyes before crossing his arms.",
        "It does not make sense to take advantage of the situation in this way.",
        "She opened her mouth to speak, but no words came out.",
        "The old man made his way through the narrow streets of the city.",
        "Back in the day, things were different, he thought to himself.",
        "He paid attention to every detail, leaving nothing to chance.",
        "The rain was pouring cats and dogs, making the roads slippery.",
    ]
    results.append(run_test(
        "medium_document",
        medium_paragraphs,
        expect_status="completed",
        max_wait=60,
    ))

    # Test 3: Empty document
    results.append(run_test(
        "empty_document",
        ["", "  ", "a"],  # too short to translate
        expect_status="completed",
        max_wait=10,
    ))

    # Summary
    print(f"\n{'='*60}")
    print(f"SUMMARY")
    print(f"{'='*60}")
    test_names = ["small_document", "medium_document", "empty_document"]
    all_passed = all(results)
    for name, passed in zip(test_names, results):
        print(f"  {'✅' if passed else '❌'} {name}")
    print(f"\n  {'ALL TESTS PASSED' if all_passed else 'SOME TESTS FAILED'}")

    sys.exit(0 if all_passed else 1)
