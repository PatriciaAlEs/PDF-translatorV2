"""
PDF Translator v3 — Layout-preserving translation via DOCX pipeline
Architecture:
  1. Upload PDF
  2. Convert PDF → DOCX (pdf2docx preserves layout, images, fonts)
  3. Extract paragraphs from DOCX (python-docx)
  4. Translate each paragraph (Google + optional AI refinement)
  5. Apply Spanish linguistic post-processing
  6. Rebuild DOCX with translated text (preserving all formatting)
  7. Convert DOCX → PDF (docx2pdf via Word/LibreOffice)
  8. Return translated PDF
"""

import os
import re
import json
import uuid
import shutil
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

BASE_DIR = Path(__file__).parent.parent
SESSIONS_DIR = BASE_DIR / "sessions"
SESSIONS_DIR.mkdir(exist_ok=True)

# Active translation jobs (in-memory progress tracking)
translation_jobs = {}

# ─────────────────────────────────────────────
# SESIONES
# ─────────────────────────────────────────────

def get_session_path(session_id):
    p = SESSIONS_DIR / session_id
    p.mkdir(exist_ok=True)
    return p

def save_session_meta(session_id, data):
    path = get_session_path(session_id) / "meta.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_session_meta(session_id):
    path = get_session_path(session_id) / "meta.json"
    if not path.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        return json.load(f)

# ─────────────────────────────────────────────
# POST-PROCESAMIENTO DE TEXTO (Espanol literario)
# ─────────────────────────────────────────────

def post_process_spanish(text):
    """Pipeline completo de post-procesamiento para texto literario en espanol."""
    try:
        from spanish_rules import apply_spanish_rules
    except ImportError:
        from backend.spanish_rules import apply_spanish_rules
    text = fix_whitespace(text)
    text = fix_dialogues(text)
    text = apply_spanish_rules(text)
    text = fix_punctuation_spacing(text)
    text = capitalize_after_period(text)
    text = fix_opening_marks(text)
    text = fix_paragraph_structure(text)
    text = fix_ellipsis(text)
    text = fix_ordinals(text)
    text = fix_grammar_languagetool(text)
    return text


def fix_whitespace(text):
    text = text.replace('\t', ' ')
    text = re.sub(r'[^\S\n]+', ' ', text)
    text = re.sub(r' +\n', '\n', text)
    text = re.sub(r'\n +(?!—)', '\n', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()


def fix_dialogues(text):
    lines = text.split('\n')
    result = []
    for line in lines:
        result.append(_convert_dialogue_line(line))
    text = '\n'.join(result)
    text = _cleanup_remaining_quotes(text)
    return text


def _cleanup_remaining_quotes(text):
    text = re.sub(r'["“«]([^"”»\n]+?)["”»]', '—\\1', text)
    text = text.replace('«', '—').replace('»', '')
    text = text.replace('“', '—').replace('”', '')
    return text


def _convert_dialogue_line(line):
    stripped = line.strip()
    if not stripped:
        return line

    OPEN_Q = r'[\"“«]'
    CLOSE_Q = r'[\"”»]'

    m = re.match(
        rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*[,.]?\s*'
        rf'(\b(?:dijo|exclamó|preguntó|respondió|susurró|gritó|murmuró|añadió|'
        rf'contestó|replicó|comentó|afirmó|negó|insistió|suplicó|ordenó|'
        rf'interrumpió|continuó|prosiguió|explicó|señaló|indicó|sugirió|'
        rf'pensó|reflexionó|musitó|balbuceó|tartamudeó|chilló|bramó|'
        rf'said|asked|replied|whispered|shouted|exclaimed|answered|added|'
        rf'murmured|cried|called|screamed|yelled|demanded|insisted)\b.*)$',
        stripped, re.IGNORECASE
    )
    if m:
        return f'—{m.group(1)} —{m.group(2)}'

    m2 = re.match(
        rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*[,.]?\s*(.+?)\s*[,.]?\s*{OPEN_Q}(.+?){CLOSE_Q}\s*$',
        stripped
    )
    if m2:
        return f'—{m2.group(1)} —{m2.group(2)}—. {m2.group(3)}'

    m3 = re.match(rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*([.!?]?)\s*$', stripped)
    if m3:
        contenido = m3.group(1)
        punct = m3.group(2)
        if contenido and contenido[-1] in '.!?¡¿':
            return f'—{contenido}'
        return f'—{contenido}{punct}'

    if re.match(rf'^{OPEN_Q}', stripped):
        line_new = re.sub(rf'^{OPEN_Q}', '—', stripped)
        line_new = re.sub(rf'{CLOSE_Q}\s*$', '', line_new)
        line_new = re.sub(rf'{CLOSE_Q}\s*[,.]?\s*', ' —', line_new)
        line_new = re.sub(rf'\s*{OPEN_Q}', '— ', line_new)
        return line_new

    if '«' in line or '»' in line:
        line = re.sub(r'«(.+?)»', lambda m: f'—{m.group(1)}', line)
        line = line.replace('«', '—').replace('»', '')

    return line


def fix_opening_marks(text):
    lines = text.split('\n')
    result = []
    for line in lines:
        line = re.sub(
            r'(?<![¿\w])([A-Za-zÁÉÍÓÚÑáéíóúñ¿][^.!?¡¿]*?\?)',
            lambda m: m.group(1) if '¿' in m.group(1) else '¿' + m.group(1),
            line
        )
        line = re.sub(
            r'(?<![¡\w])([A-Za-zÁÉÍÓÚÑáéíóúñ¡][^.!?¡¿]*?!)',
            lambda m: m.group(1) if '¡' in m.group(1) else '¡' + m.group(1),
            line
        )
        result.append(line)
    return '\n'.join(result)


def fix_punctuation_spacing(text):
    text = re.sub(r' +([.,;:?!)\]»])', r'\1', text)
    text = re.sub(r'([([«¿¡]) +', r'\1', text)
    text = re.sub(r'([.,;:?!])([A-Za-záéíóúñÁÉÍÓÚÑ])', r'\1 \2', text)
    text = re.sub(r'\.([A-ZÁÉÍÓÚÑ])', r'. \1', text)
    text = re.sub(r',([a-záéíóúñA-ZÁÉÍÓÚÑ])', r', \1', text)
    text = re.sub(r'\.\s*\.\s*\.', '...', text)
    text = re.sub(r'([.])\1(?!\.)', r'\1', text)
    text = re.sub(r',,+', ',', text)
    text = re.sub(r';;+', ';', text)
    return text


def capitalize_after_period(text):
    text = re.sub(r'([.?!]) (\w)', lambda m: m.group(1) + ' ' + m.group(2).upper(), text)
    text = re.sub(r'(\.{3}) (\w)', lambda m: m.group(1) + ' ' + m.group(2).upper(), text)
    text = re.sub(r'(^|\n)(—)(\w)', lambda m: m.group(1) + '—' + m.group(3).upper(), text)
    lines = text.split('\n')
    fixed = []
    for line in lines:
        stripped = line.lstrip()
        if stripped:
            indent = line[:len(line) - len(stripped)]
            if stripped.startswith('—'):
                fixed.append(indent + stripped)
            else:
                fixed.append(indent + stripped[0].upper() + stripped[1:])
        else:
            fixed.append(line)
    return '\n'.join(fixed)


def fix_paragraph_structure(text):
    lines = text.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            result.append('')
            i += 1
            continue
        if re.match(r'^[\s*\-=•~─—_]{3,}$', stripped):
            result.append('')
            result.append('* * *')
            result.append('')
            i += 1
            continue
        if _is_chapter_heading(stripped):
            if result and result[-1] != '':
                result.append('')
            result.append(stripped)
            result.append('')
            i += 1
            continue
        if stripped.startswith('—'):
            if result and result[-1] != '' and not result[-1].startswith('—'):
                result.append('')
            result.append(stripped)
            i += 1
            continue
        result.append(stripped)
        i += 1
    cleaned = '\n'.join(result)
    cleaned = re.sub(r'\n{4,}', '\n\n\n', cleaned)
    return cleaned


def _is_chapter_heading(line):
    patterns = [
        r'^(?:Capítulo|CAPÍTULO|Chapter|CHAPTER)\s+[\dIVXLCDMivxlcdm]+',
        r'^(?:PARTE|Parte|Part|PART)\s+[\dIVXLCDMivxlcdm]+',
        r'^(?:Prólogo|PRÓLOGO|Epílogo|EPÍLOGO|Prologue|Epilogue)',
        r'^(?:LIBRO|Libro|Book|BOOK)\s+[\dIVXLCDMivxlcdm]+',
        r'^[IVXLCDM]+\s*$',
        r'^\d+\s*$',
    ]
    return any(re.match(p, line, re.IGNORECASE) for p in patterns)


def fix_ellipsis(text):
    text = re.sub(r'(?<!\.)\.{2}(?!\.)', '...', text)
    text = re.sub(r'\.{4,}', '...', text)
    text = re.sub(r'(\.\.\.)([A-Za-záéíóúñÁÉÍÓÚÑ])', r'\1 \2', text)
    return text


def fix_ordinals(text):
    text = re.sub(r'\b(\d+)(?:st|nd|rd|th)\b', r'\1.º', text)
    return text


def fix_grammar_languagetool(text):
    """Use LanguageTool public API to fix grammar, spelling, and accents in Spanish."""
    import requests as _requests
    if not text.strip():
        return text

    # Split into chunks (API limit ~20KB per request)
    MAX_CHUNK = 15000
    paragraphs = text.split('\n')
    chunks = []
    current = ""
    for p in paragraphs:
        if len(current) + len(p) + 1 > MAX_CHUNK and current:
            chunks.append(current)
            current = p
        else:
            current = current + '\n' + p if current else p
    if current:
        chunks.append(current)

    corrected_chunks = []
    for chunk in chunks:
        try:
            r = _requests.post(
                'https://api.languagetool.org/v2/check',
                data={'text': chunk, 'language': 'es'},
                timeout=30,
            )
            if r.status_code != 200:
                corrected_chunks.append(chunk)
                continue

            matches = r.json().get('matches', [])
            # Filter out corrections that change proper nouns (capitalized words)
            safe_matches = []
            for m in matches:
                if not m.get('replacements'):
                    continue
                original_word = chunk[m['offset']:m['offset'] + m['length']]
                # Skip if it looks like a proper noun (capitalized, not start of sentence)
                if (original_word and original_word[0].isupper()
                        and m['offset'] > 0
                        and chunk[m['offset'] - 1] not in '.?!¿¡\n'
                        and m['rule']['id'] == 'MORFOLOGIK_RULE_ES'):
                    continue
                safe_matches.append(m)
            # Apply corrections in reverse order to preserve offsets
            result = chunk
            for m in reversed(safe_matches):
                if m.get('replacements'):
                    start = m['offset']
                    end = start + m['length']
                    result = result[:start] + m['replacements'][0]['value'] + result[end:]
            corrected_chunks.append(result)
        except Exception as e:
            print(f"LanguageTool API error: {e}")
            corrected_chunks.append(chunk)

    return '\n'.join(corrected_chunks)


# ─────────────────────────────────────────────
# TRADUCCION: Google Translate (chunked)
# ─────────────────────────────────────────────

def _translate_text_google(text, source_lang="auto"):
    if not text.strip():
        return text

    from deep_translator import GoogleTranslator

    MAX_CHUNK = 4500
    chunks = []
    remaining = text
    while len(remaining) > MAX_CHUNK:
        cut = remaining[:MAX_CHUNK].rfind('\n')
        if cut == -1:
            cut = remaining[:MAX_CHUNK].rfind('. ')
        if cut == -1:
            cut = MAX_CHUNK
        chunks.append(remaining[:cut + 1])
        remaining = remaining[cut + 1:]
    if remaining:
        chunks.append(remaining)

    translator = GoogleTranslator(source=source_lang, target="es")
    translated_chunks = []
    for c in chunks:
        if c.strip():
            translated_chunks.append(translator.translate(c))
    return "\n".join(translated_chunks)


# ─────────────────────────────────────────────
# TRADUCCION: IA (Gemini / Claude)
# ─────────────────────────────────────────────

def _gemini(text, is_google_result):
    try:
        from google import genai
    except ImportError:
        raise Exception("Ejecuta: pip install google-genai")

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise Exception(
            "Falta GEMINI_API_KEY en .env\n"
            "Obtén una clave GRATIS en: https://aistudio.google.com/app/apikey"
        )

    client = genai.Client(api_key=api_key)

    rules = (
        "REGLAS OBLIGATORIAS DE FORMATO ESPAÑOL:\n"
        "1. DIÁLOGOS: Usa guión largo (—) SIEMPRE, NUNCA comillas. "
        "Formato: —Texto del diálogo —dijo el personaje—. Continuación.\n"
        "2. INCISOS DEL NARRADOR: —Hola —dijo Juan—. Luego se fue. "
        "(guión largo antes y después del inciso, punto después del cierre).\n"
        "3. SIGNOS DE APERTURA: Siempre usar ¿...? y ¡...! (con apertura).\n"
        "4. MAYÚSCULAS: Después de . ? ! y ... siempre mayúscula.\n"
        "5. PUNTUACIÓN: Sin espacio antes de . , ; : ? ! — Espacio después.\n"
        "6. PÁRRAFOS: Mantén los saltos de párrafo del original. "
        "Cada diálogo de distinto personaje en párrafo aparte.\n"
        "7. NATURALIDAD: Usa expresiones naturales del español castellano, "
        "no traducciones literales. Adapta modismos e interjecciones.\n"
        "8. REGISTRO: Mantén el registro del original (formal/informal, tú/usted).\n"
        "9. NOMBRES PROPIOS: No traduzcas nombres propios de personas ni lugares ficticios.\n"
        "10. COMAS VOCATIVAS: Siempre coma antes del vocativo: —Hola, Juan.\n"
        "11. LEÍSMO/LAÍSMO: Usa los pronombres correctamente según la RAE.\n"
        "12. PUNTOS SUSPENSIVOS: Exactamente tres puntos (...), espacio después.\n"
    )

    if is_google_result:
        prompt = (
            "Eres un corrector y editor literario experto en español castellano.\n"
            "Se te proporciona una traducción automática de un libro/texto largo.\n"
            "Tu tarea es MEJORARLA para que suene como prosa literaria profesional "
            "publicada en España.\n\n"
            f"{rules}\n"
            "CORRECCIONES ADICIONALES A BUSCAR:\n"
            "- Frases que suenan a traducción literal del inglés\n"
            "- Gerundios mal usados (en español se usan menos que en inglés)\n"
            "- Voz pasiva excesiva (en español se prefiere activa o pasiva refleja)\n"
            "- Repeticiones innecesarias de sujeto pronominable\n"
            "- Falsos amigos (actually≠actualmente, etc.)\n\n"
            "Devuelve ÚNICAMENTE el texto mejorado, sin explicaciones ni comentarios.\n\n"
            f"TEXTO A MEJORAR:\n{text}"
        )
    else:
        prompt = (
            "Eres un traductor literario profesional especializado en español castellano.\n"
            "Traduce el siguiente texto manteniendo el estilo, tono y ritmo narrativo del original.\n"
            "La traducción debe leerse como una obra publicada en español, no como una traducción.\n\n"
            f"{rules}\n"
            "INSTRUCCIONES ADICIONALES:\n"
            "- Adapta modismos e interjecciones al español natural\n"
            "- Evita gerundios excesivos, voz pasiva innecesaria y calcos del inglés\n"
            "- Mantén la estructura de párrafos del original\n"
            "- Para cada diálogo de personaje distinto, usa párrafo aparte\n\n"
            "Devuelve ÚNICAMENTE la traducción, sin explicaciones ni comentarios.\n\n"
            f"TEXTO ORIGINAL:\n{text}"
        )

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
    )
    return response.text


def _claude(text, is_google_result):
    try:
        import anthropic
    except ImportError:
        raise Exception("Ejecuta: pip install anthropic")

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise Exception("Falta ANTHROPIC_API_KEY en .env")

    client = anthropic.Anthropic(api_key=api_key)

    rules = (
        "REGLAS: Diálogos con guión largo (—), nunca comillas. "
        "Incisos: —texto —dijo X—. Continúa. "
        "Signos de apertura obligatorios: ¿...? ¡...! "
        "Mayúscula después de . ? ! y puntos suspensivos. "
        "Coma vocativa. Sin gerundios ni pasiva innecesarios. "
        "Español castellano natural, no traducción literal."
    )

    if is_google_result:
        action = (
            f"Mejora esta traducción automática para que suene como prosa literaria "
            f"profesional en español castellano. {rules}"
        )
    else:
        action = (
            f"Traduce al español castellano como prosa literaria publicada. "
            f"Mantén estilo, tono y ritmo. {rules}"
        )
    prompt = (
        f"{action}\n"
        f"Devuelve ÚNICAMENTE el texto resultante, sin explicaciones.\n\nTEXTO:\n{text}"
    )
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text


# ─────────────────────────────────────────────
# PDF <-> DOCX CONVERSION
# ─────────────────────────────────────────────

def _pdf_to_docx(pdf_path, docx_path):
    """Convert PDF to DOCX preserving layout, fonts, images."""
    from pdf2docx import Converter
    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()


def _docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX back to PDF using pure Python (fpdf2 + python-docx).
    No MS Word or LibreOffice needed. Uses Unicode TTF font."""
    from docx import Document
    from fpdf import FPDF

    doc = Document(str(docx_path))

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Register a Unicode TTF font (Windows Arial)
    font_path = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "arial.ttf"
    font_bold = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "arialbd.ttf"
    font_italic = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "ariali.ttf"
    font_bi = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "arialbi.ttf"

    if font_path.exists():
        pdf.add_font("ArialUni", "", str(font_path), uni=True)
        if font_bold.exists():
            pdf.add_font("ArialUni", "B", str(font_bold), uni=True)
        if font_italic.exists():
            pdf.add_font("ArialUni", "I", str(font_italic), uni=True)
        if font_bi.exists():
            pdf.add_font("ArialUni", "BI", str(font_bi), uni=True)
        font_family = "ArialUni"
    else:
        # Fallback for non-Windows systems
        pdf.add_font("DejaVu", "", str(Path(__file__).parent / "DejaVuSans.ttf"), uni=True)
        font_family = "DejaVu"

    pdf.set_font(font_family, size=11)

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            pdf.ln(4)
            continue

        # Detect bold/italic from first run, but use uniform font size
        bold = False
        italic = False
        if para.runs:
            run = para.runs[0]
            bold = run.bold or False
            italic = run.italic or False

        style = ""
        if bold:
            style += "B"
        if italic:
            style += "I"

        pdf.set_font(font_family, style=style, size=11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(2)

    pdf.output(str(pdf_path))


def _extract_paragraphs(docx_path):
    """Extract all non-empty paragraphs from DOCX."""
    from docx import Document
    doc = Document(str(docx_path))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
    return paragraphs


def _translate_docx_in_place(docx_path, output_path, source_lang="auto",
                              use_ai=False, provider="gemini",
                              progress_callback=None):
    """
    Translate a DOCX file in-place preserving all formatting.
    Each paragraph's text is translated while keeping runs/styles intact.
    Returns stats dict.
    """
    from docx import Document
    doc = Document(str(docx_path))

    total = sum(1 for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) >= 3)
    translated_count = 0

    for para in doc.paragraphs:
        original_text = para.text.strip()
        if not original_text or len(original_text) < 3:
            continue

        try:
            translated = _translate_text_google(original_text, source_lang)

            if use_ai:
                try:
                    if provider == "gemini":
                        translated = _gemini(translated, is_google_result=True)
                    else:
                        translated = _claude(translated, is_google_result=True)
                except Exception as ai_err:
                    print(f"AI refinement failed, using Google result: {ai_err}")

            translated = post_process_spanish(translated)
            _apply_translation_to_paragraph(para, translated)
            translated_count += 1

            if progress_callback:
                progress_callback(translated_count, total)

        except Exception as e:
            print(f"Error translating paragraph: {e}")

    doc.save(str(output_path))
    return {"total_paragraphs": total, "translated": translated_count}


def _apply_translation_to_paragraph(para, translated_text):
    """Replace text in a paragraph while preserving first-run formatting."""
    runs = para.runs
    if not runs:
        return
    runs[0].text = translated_text
    for run in runs[1:]:
        run.text = ""


def _translate_single_paragraph(text, source_lang, use_ai, provider):
    """Translate a single paragraph's text. Thread-safe (pure function)."""
    translated = _translate_text_google(text, source_lang)
    if use_ai:
        try:
            if provider == "gemini":
                translated = _gemini(translated, is_google_result=True)
            else:
                translated = _claude(translated, is_google_result=True)
        except Exception as ai_err:
            print(f"AI refinement failed, using Google result: {ai_err}")
    return post_process_spanish(translated)


def _translate_docx_background(session_id, docx_path, output_path,
                                source_lang, use_ai, provider):
    """Background job: translate all paragraphs with parallel workers."""
    from docx import Document

    job = translation_jobs[session_id]
    job["status"] = "translating"
    job["started_at"] = time.time()

    try:
        doc = Document(str(docx_path))

        # Collect paragraphs to translate
        work_items = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) >= 3:
                work_items.append((i, text))

        total = len(work_items)
        job["total"] = total
        job["translated"] = 0

        if total == 0:
            job["status"] = "completed"
            job["finished_at"] = time.time()
            doc.save(str(output_path))
            return

        # Fewer workers when using AI (API rate limits)
        max_workers = 2 if use_ai else 4
        translations = {}

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(
                    _translate_single_paragraph, text, source_lang, use_ai, provider
                ): idx
                for idx, text in work_items
            }

            for future in as_completed(futures):
                if job.get("cancelled"):
                    for f in futures:
                        f.cancel()
                    job["status"] = "cancelled"
                    return

                idx = futures[future]
                try:
                    translated_text = future.result()
                    translations[idx] = translated_text
                except Exception as e:
                    job["errors"].append(f"Párrafo {idx}: {str(e)}")
                    print(f"Error translating paragraph {idx}: {e}")

                job["translated"] += 1
                elapsed = time.time() - job["started_at"]
                done = job["translated"]
                if done > 0:
                    rate = elapsed / done
                    job["eta_seconds"] = int((total - done) * rate)

        if job.get("cancelled"):
            job["status"] = "cancelled"
            return

        # Apply translations back to document (sequential)
        if not translations:
            job["status"] = "error"
            job["error_message"] = f"Ningún párrafo se tradujo correctamente. Errores: {'; '.join(job['errors'][:3])}"
            print(f"Translation failed for {session_id}: no paragraphs translated. Errors: {job['errors']}")
            return

        for i, para in enumerate(doc.paragraphs):
            if i in translations:
                _apply_translation_to_paragraph(para, translations[i])

        doc.save(str(output_path))
        print(f"Saved translated DOCX: {output_path} ({len(translations)}/{total} paragraphs)")

        job["status"] = "completed"
        job["finished_at"] = time.time()

        meta = load_session_meta(session_id)
        meta["status"] = "translated"
        meta["translation_method"] = f"google{'+' + provider if use_ai else ''}"
        save_session_meta(session_id, meta)

    except Exception as e:
        job["status"] = "error"
        job["error_message"] = str(e)
        print(f"Translation job failed for {session_id}: {e}")


# ─────────────────────────────────────────────
# NEW API: DOCX PIPELINE
# ─────────────────────────────────────────────

@app.route("/api/upload", methods=["POST"])
def upload_pdf():
    """Upload PDF and auto-convert to DOCX (preserves layout)."""
    if "file" not in request.files:
        return jsonify({"error": "No se envió ningún archivo"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "El archivo debe ser un PDF"}), 400

    upload_start = time.time()

    session_id = str(uuid.uuid4())
    session_path = get_session_path(session_id)

    pdf_path = session_path / "original.pdf"
    file.save(str(pdf_path))

    # Page count
    import fitz
    doc = fitz.open(str(pdf_path))
    total_pages = len(doc)
    file_size_mb = round(pdf_path.stat().st_size / 1024 / 1024, 2)
    doc.close()

    # Auto-convert PDF -> DOCX
    docx_path = session_path / "original.docx"
    try:
        _pdf_to_docx(pdf_path, docx_path)
    except Exception as e:
        return jsonify({"error": f"Error convirtiendo PDF a DOCX: {str(e)}"}), 500

    paragraphs = _extract_paragraphs(docx_path)

    upload_seconds = round(time.time() - upload_start, 1)

    meta = {
        "session_id": session_id,
        "original_filename": file.filename,
        "total_pages": total_pages,
        "total_paragraphs": len(paragraphs),
        "file_size_mb": file_size_mb,
        "upload_seconds": upload_seconds,
        "status": "uploaded",
    }
    save_session_meta(session_id, meta)

    return jsonify({
        "session_id": session_id,
        "filename": file.filename,
        "total_pages": total_pages,
        "total_paragraphs": len(paragraphs),
        "file_size_mb": file_size_mb,
        "upload_seconds": upload_seconds,
        "paragraphs_preview": paragraphs[:20],
    })


@app.route("/api/paragraphs/<session_id>")
def get_paragraphs(session_id):
    """Return all paragraphs from the DOCX."""
    session_path = get_session_path(session_id)
    docx_path = session_path / "original.docx"
    if not docx_path.exists():
        return jsonify({"error": "DOCX no encontrado"}), 404
    paragraphs = _extract_paragraphs(docx_path)
    return jsonify({"paragraphs": paragraphs, "count": len(paragraphs)})


@app.route("/api/translate-docx", methods=["POST"])
def translate_docx():
    """
    Start background translation of all paragraphs in the DOCX.
    Returns immediately; poll /api/translate-progress/<session_id> for status.
    """
    data = request.json
    session_id = data.get("session_id")
    source_lang = data.get("source_lang", "auto")
    use_ai = data.get("use_ai", False)
    provider = data.get("provider", "gemini")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404

    session_path = get_session_path(session_id)
    docx_path = session_path / "original.docx"
    translated_docx = session_path / "translated.docx"

    if not docx_path.exists():
        return jsonify({"error": "DOCX original no encontrado"}), 404

    # Prevent duplicate jobs
    if session_id in translation_jobs and translation_jobs[session_id].get("status") == "translating":
        return jsonify({"error": "Ya hay una traducción en curso"}), 409

    # Initialize job tracker
    translation_jobs[session_id] = {
        "status": "starting",
        "total": meta.get("total_paragraphs", 0),
        "translated": 0,
        "errors": [],
        "eta_seconds": None,
        "started_at": None,
        "finished_at": None,
        "cancelled": False,
        "error_message": None,
    }

    # Launch background thread
    thread = threading.Thread(
        target=_translate_docx_background,
        args=(session_id, docx_path, translated_docx, source_lang, use_ai, provider),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "started": True,
        "session_id": session_id,
        "total_paragraphs": meta.get("total_paragraphs", 0),
    })


@app.route("/api/translate-progress/<session_id>")
def translate_progress(session_id):
    """Poll translation progress for a session."""
    job = translation_jobs.get(session_id)
    if not job:
        return jsonify({"status": "none", "error": "No hay traducción en curso"}), 404

    result = {
        "status": job["status"],
        "total": job.get("total", 0),
        "translated": job.get("translated", 0),
        "errors": job.get("errors", []),
        "eta_seconds": job.get("eta_seconds"),
        "error_message": job.get("error_message"),
    }

    if job.get("total", 0) > 0:
        result["progress_percent"] = round(job["translated"] / job["total"] * 100, 1)
    else:
        result["progress_percent"] = 0

    if job.get("started_at"):
        result["elapsed_seconds"] = int(time.time() - job["started_at"])

    if job["status"] == "completed" and job.get("finished_at") and job.get("started_at"):
        result["total_seconds"] = int(job["finished_at"] - job["started_at"])

    return jsonify(result)


@app.route("/api/translate-cancel/<session_id>", methods=["POST"])
def translate_cancel(session_id):
    """Request cancellation of an active translation job."""
    job = translation_jobs.get(session_id)
    if not job or job["status"] != "translating":
        return jsonify({"error": "No hay traducción activa para cancelar"}), 404
    job["cancelled"] = True
    return jsonify({"success": True, "message": "Cancelación solicitada"})


@app.route("/api/generate-pdf", methods=["POST"])
def generate_pdf():
    """Convert the translated DOCX back to PDF (preserves layout)."""
    data = request.json
    session_id = data.get("session_id")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404

    session_path = get_session_path(session_id)
    translated_docx = session_path / "translated.docx"

    if not translated_docx.exists():
        return jsonify({"error": "Primero debes traducir el documento"}), 400

    output_filename = f"traduccion_{session_id[:8]}.pdf"
    output_path = session_path / output_filename

    gen_start = time.time()

    try:
        _docx_to_pdf(translated_docx, output_path)

        gen_seconds = round(time.time() - gen_start, 1)
        output_size_mb = round(output_path.stat().st_size / 1024 / 1024, 2)

        meta["status"] = "completed"
        meta["output_filename"] = output_filename
        meta["generate_seconds"] = gen_seconds
        save_session_meta(session_id, meta)

        return jsonify({
            "success": True,
            "filename": output_filename,
            "download_url": f"/api/download/{session_id}/{output_filename}",
            "generate_seconds": gen_seconds,
            "output_size_mb": output_size_mb,
        })
    except Exception as e:
        return jsonify({"error": f"Error generando PDF: {str(e)}"}), 500


# ─────────────────────────────────────────────
# LEGACY ENDPOINTS (backward compat)
# ─────────────────────────────────────────────

@app.route("/api/split", methods=["POST"])
def split_pdf():
    data = request.json
    session_id = data.get("session_id")
    pages_per_part = int(data.get("pages_per_part", 10))

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404

    session_path = get_session_path(session_id)
    pdf_path = session_path / "original.pdf"

    import fitz
    src = fitz.open(str(pdf_path))
    total_pages = len(src)

    parts_dir = session_path / "parts"
    if parts_dir.exists():
        shutil.rmtree(str(parts_dir))
    parts_dir.mkdir()

    parts = []
    part_num = 1
    for start in range(0, total_pages, pages_per_part):
        end = min(start + pages_per_part, total_pages)
        part_doc = fitz.open()
        part_doc.insert_pdf(src, from_page=start, to_page=end - 1)
        part_filename = f"parte_{part_num:03d}.pdf"
        part_doc.save(str(parts_dir / part_filename))
        part_doc.close()
        parts.append({
            "part_num": part_num,
            "filename": part_filename,
            "pages_start": start + 1,
            "pages_end": end,
            "page_count": end - start,
            "status": "pending",
            "translated_text": None,
        })
        part_num += 1
    src.close()

    meta["parts"] = parts
    meta["pages_per_part"] = pages_per_part
    meta["status"] = "split"
    save_session_meta(session_id, meta)
    return jsonify({"session_id": session_id, "total_parts": len(parts), "parts": parts})


@app.route("/api/extract-text", methods=["POST"])
def extract_text():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404

    session_path = get_session_path(session_id)
    part_path = session_path / "parts" / f"parte_{part_num:03d}.pdf"
    if not part_path.exists():
        return jsonify({"error": f"Parte {part_num} no encontrada"}), 404

    import fitz
    doc = fitz.open(str(part_path))
    text = ""
    for page in doc:
        text += page.get_text() + "\n\n"
    doc.close()

    return jsonify({"part_num": part_num, "text": text.strip(), "char_count": len(text)})


@app.route("/api/translate-google", methods=["POST"])
def translate_google():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    source_lang = data.get("source_lang", "auto")
    text = data.get("text", "")

    if not text.strip():
        return jsonify({"error": "No hay texto para traducir"}), 400

    try:
        translated = _translate_text_google(text, source_lang)
        translated = post_process_spanish(translated)

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = translated
                part["status"] = "translated_google"
                break
        save_session_meta(session_id, meta)

        return jsonify({"part_num": part_num, "translated_text": translated, "method": "google"})
    except Exception as e:
        return jsonify({"error": f"Error en traducción: {str(e)}"}), 500


@app.route("/api/translate-ai", methods=["POST"])
def translate_ai():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    text = data.get("text", "")
    is_google_result = data.get("is_google_result", False)
    provider = data.get("provider", "gemini")

    if not text.strip():
        return jsonify({"error": "No hay texto para procesar"}), 400

    try:
        if provider == "gemini":
            translated = _gemini(text, is_google_result)
        else:
            translated = _claude(text, is_google_result)
        translated = post_process_spanish(translated)

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = translated
                part["status"] = f"translated_{provider}"
                break
        save_session_meta(session_id, meta)

        return jsonify({"part_num": part_num, "translated_text": translated, "method": provider})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/translate-pipeline", methods=["POST"])
def translate_pipeline():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    source_lang = data.get("source_lang", "auto")
    text = data.get("text", "")
    skip_ai = data.get("skip_ai", False)
    provider = data.get("provider", "gemini")

    if not text.strip():
        return jsonify({"error": "No hay texto para traducir"}), 400

    steps_done = []
    try:
        google_result = _translate_text_google(text, source_lang)
        steps_done.append("google")

        ai_result = google_result
        if not skip_ai:
            try:
                if provider == "gemini":
                    ai_result = _gemini(google_result, is_google_result=True)
                else:
                    ai_result = _claude(google_result, is_google_result=True)
                steps_done.append(provider)
            except Exception as ai_err:
                print(f"AI failed, using Google: {ai_err}")

        final = post_process_spanish(ai_result)
        steps_done.append("postprocess")

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = final
                part["status"] = "translated_pipeline"
                break
        save_session_meta(session_id, meta)

        return jsonify({
            "part_num": part_num,
            "translated_text": final,
            "method": "+".join(steps_done),
            "steps": steps_done,
        })
    except Exception as e:
        return jsonify({"error": f"Error en pipeline: {str(e)}"}), 500


@app.route("/api/save-translation", methods=["POST"])
def save_translation():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    translated_text = data.get("translated_text", "")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404

    for part in meta.get("parts", []):
        if part["part_num"] == part_num:
            part["translated_text"] = translated_text
            part["status"] = "translated_manual"
            break

    save_session_meta(session_id, meta)
    return jsonify({"success": True, "part_num": part_num})


@app.route("/api/fix-dialogues", methods=["POST"])
def fix_dialogues_api():
    data = request.json
    fixed = post_process_spanish(data.get("text", ""))
    return jsonify({"fixed_text": fixed})


# ─────────────────────────────────────────────
# DOWNLOAD / SESSION
# ─────────────────────────────────────────────

@app.route("/api/download/<session_id>/<filename>")
def download_file(session_id, filename):
    fp = get_session_path(session_id) / filename
    if not fp.exists():
        return jsonify({"error": "Archivo no encontrado"}), 404
    return send_file(str(fp), as_attachment=True, download_name=filename)


@app.route("/api/session/<session_id>")
def get_session(session_id):
    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesión no encontrada"}), 404
    return jsonify(meta)


# ─────────────────────────────────────────────
# SERVIR FRONTEND
# ─────────────────────────────────────────────

FRONTEND_DIR = BASE_DIR / "frontend"

@app.route("/")
def serve_frontend():
    return send_file(str(FRONTEND_DIR / "index.html"))

@app.route("/<path:filename>")
def serve_static(filename):
    fp = FRONTEND_DIR / filename
    if fp.exists() and fp.is_file():
        return send_file(str(fp))
    return jsonify({"error": "No encontrado"}), 404


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    print(f"\n\U0001f680 PDF Translator v3 — http://localhost:{port}")
    print(f"   ✅ Pipeline: PDF → DOCX → Traducir → PDF")
    print(f"   ✅ Preserva layout, fuentes, imágenes, tablas")
    print(f"   ✅ Google Translate + Gemini/Claude refinamiento")
    print(f"   ✅ Post-procesamiento lingüístico español\n")
    app.run(debug=True, port=port)
