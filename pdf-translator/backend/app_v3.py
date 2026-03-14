"""
PDF Translator v3 — Layout-preserving translation via DOCX pipeline
Architecture:
  1. Upload PDF
  2. Convert PDF → DOCX (pdf2docx preserves layout, images, fonts)
  3. Extract paragraphs from DOCX (python-docx)
  4. Translate each paragraph (Google + optional AI refinement)
  5. Apply Spanish linguistic post-processing
  6. Rebuild DOCX with translated text (preserving all formatting)
  7. Convert DOCX → PDF (docx2pdf via Word/LibreOffice)
  8. Return translated PDF
"""

import os
import re
import json
import uuid
import shutil
from pathlib import Path
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

BASE_DIR = Path(__file__).parent.parent
SESSIONS_DIR = BASE_DIR / "sessions"
SESSIONS_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────
# SESIONES
# ─────────────────────────────────────────────

def get_session_path(session_id):
    p = SESSIONS_DIR / session_id
    p.mkdir(exist_ok=True)
    return p

def save_session_meta(session_id, data):
    path = get_session_path(session_id) / "meta.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_session_meta(session_id):
    path = get_session_path(session_id) / "meta.json"
    if not path.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        return json.load(f)

# ─────────────────────────────────────────────
# POST-PROCESAMIENTO DE TEXTO (Espanol literario)
# ─────────────────────────────────────────────

def post_process_spanish(text):
    """Pipeline completo de post-procesamiento para texto literario en espanol."""
    text = fix_whitespace(text)
    text = fix_dialogues(text)
    text = fix_punctuation_spacing(text)
    text = capitalize_after_period(text)
    text = fix_opening_marks(text)
    text = fix_paragraph_structure(text)
    text = fix_ellipsis(text)
    text = fix_ordinals(text)
    return text


def fix_whitespace(text):
    text = text.replace('\t', ' ')
    text = re.sub(r'[^\S\n]+', ' ', text)
    text = re.sub(r' +\n', '\n', text)
    text = re.sub(r'\n +(?!\u2014)', '\n', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()


def fix_dialogues(text):
    lines = text.split('\n')
    result = []
    for line in lines:
        result.append(_convert_dialogue_line(line))
    text = '\n'.join(result)
    text = _cleanup_remaining_quotes(text)
    return text


def _cleanup_remaining_quotes(text):
    text = re.sub(r'["\u201c\u00ab]([^"\u201d\u00bb\n]+?)["\u201d\u00bb]', r'\u2014\1', text)
    text = text.replace('\u00ab', '\u2014').replace('\u00bb', '')
    text = text.replace('\u201c', '\u2014').replace('\u201d', '')
    return text


def _convert_dialogue_line(line):
    stripped = line.strip()
    if not stripped:
        return line

    OPEN_Q = r'[\"\u201c\u00ab]'
    CLOSE_Q = r'[\"\u201d\u00bb]'

    m = re.match(
        rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*[,.]?\s*'
        rf'(\b(?:dijo|exclam\u00f3|pregunt\u00f3|respondi\u00f3|susurr\u00f3|grit\u00f3|murmur\u00f3|a\u00f1adi\u00f3|'
        rf'contest\u00f3|replic\u00f3|coment\u00f3|afirm\u00f3|neg\u00f3|insisti\u00f3|suplic\u00f3|orden\u00f3|'
        rf'interrumpi\u00f3|continu\u00f3|prosigui\u00f3|explic\u00f3|se\u00f1al\u00f3|indic\u00f3|sugiri\u00f3|'
        rf'pens\u00f3|reflexion\u00f3|musit\u00f3|balbuce\u00f3|tartamude\u00f3|chill\u00f3|bram\u00f3|'
        rf'said|asked|replied|whispered|shouted|exclaimed|answered|added|'
        rf'murmured|cried|called|screamed|yelled|demanded|insisted)\b.*)$',
        stripped, re.IGNORECASE
    )
    if m:
        return f'\u2014{m.group(1)} \u2014{m.group(2)}'

    m2 = re.match(
        rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*[,.]?\s*(.+?)\s*[,.]?\s*{OPEN_Q}(.+?){CLOSE_Q}\s*$',
        stripped
    )
    if m2:
        return f'\u2014{m2.group(1)} \u2014{m2.group(2)}\u2014. {m2.group(3)}'

    m3 = re.match(rf'^{OPEN_Q}(.+?){CLOSE_Q}\s*([.!?]?)\s*$', stripped)
    if m3:
        contenido = m3.group(1)
        punct = m3.group(2)
        if contenido and contenido[-1] in '.!?\u00a1\u00bf':
            return f'\u2014{contenido}'
        return f'\u2014{contenido}{punct}'

    if re.match(rf'^{OPEN_Q}', stripped):
        line_new = re.sub(rf'^{OPEN_Q}', '\u2014', stripped)
        line_new = re.sub(rf'{CLOSE_Q}\s*$', '', line_new)
        line_new = re.sub(rf'{CLOSE_Q}\s*[,.]?\s*', ' \u2014', line_new)
        line_new = re.sub(rf'\s*{OPEN_Q}', '\u2014 ', line_new)
        return line_new

    if '\u00ab' in line or '\u00bb' in line:
        line = re.sub(r'\u00ab(.+?)\u00bb', lambda m: f'\u2014{m.group(1)}', line)
        line = line.replace('\u00ab', '\u2014').replace('\u00bb', '')

    return line


def fix_opening_marks(text):
    lines = text.split('\n')
    result = []
    for line in lines:
        line = re.sub(
            r'(?<![\u00bf\w])([A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00bf][^.!?\u00a1\u00bf]*?\?)',
            lambda m: m.group(1) if '\u00bf' in m.group(1) else '\u00bf' + m.group(1),
            line
        )
        line = re.sub(
            r'(?<![\u00a1\w])([A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00a1][^.!?\u00a1\u00bf]*?!)',
            lambda m: m.group(1) if '\u00a1' in m.group(1) else '\u00a1' + m.group(1),
            line
        )
        result.append(line)
    return '\n'.join(result)


def fix_punctuation_spacing(text):
    text = re.sub(r' +([.,;:?!)\]\u00bb])', r'\1', text)
    text = re.sub(r'([([\u00ab\u00bf\u00a1]) +', r'\1', text)
    text = re.sub(r'([.,;:?!])([A-Za-z\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00c1\u00c9\u00cd\u00d3\u00da\u00d1])', r'\1 \2', text)
    text = re.sub(r'\.([A-Z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1])', r'. \1', text)
    text = re.sub(r',([a-z\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1A-Z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1])', r', \1', text)
    text = re.sub(r'\.\s*\.\s*\.', '...', text)
    text = re.sub(r'([.])\1(?!\.)', r'\1', text)
    text = re.sub(r',,+', ',', text)
    text = re.sub(r';;+', ';', text)
    return text


def capitalize_after_period(text):
    text = re.sub(r'([.?!]) (\w)', lambda m: m.group(1) + ' ' + m.group(2).upper(), text)
    text = re.sub(r'(\.{3}) (\w)', lambda m: m.group(1) + ' ' + m.group(2).upper(), text)
    text = re.sub(r'(^|\n)(\u2014)(\w)', lambda m: m.group(1) + '\u2014' + m.group(3).upper(), text)
    lines = text.split('\n')
    fixed = []
    for line in lines:
        stripped = line.lstrip()
        if stripped:
            indent = line[:len(line) - len(stripped)]
            if stripped.startswith('\u2014'):
                fixed.append(indent + stripped)
            else:
                fixed.append(indent + stripped[0].upper() + stripped[1:])
        else:
            fixed.append(line)
    return '\n'.join(fixed)


def fix_paragraph_structure(text):
    lines = text.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            result.append('')
            i += 1
            continue
        if re.match(r'^[\s*\-=\u2022~\u2500\u2014_]{3,}$', stripped):
            result.append('')
            result.append('* * *')
            result.append('')
            i += 1
            continue
        if _is_chapter_heading(stripped):
            if result and result[-1] != '':
                result.append('')
            result.append(stripped)
            result.append('')
            i += 1
            continue
        if stripped.startswith('\u2014'):
            if result and result[-1] != '' and not result[-1].startswith('\u2014'):
                result.append('')
            result.append(stripped)
            i += 1
            continue
        result.append(stripped)
        i += 1
    cleaned = '\n'.join(result)
    cleaned = re.sub(r'\n{4,}', '\n\n\n', cleaned)
    return cleaned


def _is_chapter_heading(line):
    patterns = [
        r'^(?:Cap\u00edtulo|CAP\u00cdTULO|Chapter|CHAPTER)\s+[\dIVXLCDMivxlcdm]+',
        r'^(?:PARTE|Parte|Part|PART)\s+[\dIVXLCDMivxlcdm]+',
        r'^(?:Pr\u00f3logo|PR\u00d3LOGO|Ep\u00edlogo|EP\u00cdLOGO|Prologue|Epilogue)',
        r'^(?:LIBRO|Libro|Book|BOOK)\s+[\dIVXLCDMivxlcdm]+',
        r'^[IVXLCDM]+\s*$',
        r'^\d+\s*$',
    ]
    return any(re.match(p, line, re.IGNORECASE) for p in patterns)


def fix_ellipsis(text):
    text = re.sub(r'(?<!\.)\.{2}(?!\.)', '...', text)
    text = re.sub(r'\.{4,}', '...', text)
    text = re.sub(r'(\.\.\.)([A-Za-z\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00c1\u00c9\u00cd\u00d3\u00da\u00d1])', r'\1 \2', text)
    return text


def fix_ordinals(text):
    text = re.sub(r'\b(\d+)(?:st|nd|rd|th)\b', r'\1.\u00ba', text)
    return text


# ─────────────────────────────────────────────
# TRADUCCION: Google Translate (chunked)
# ─────────────────────────────────────────────

def _translate_text_google(text, source_lang="auto"):
    if not text.strip():
        return text

    from deep_translator import GoogleTranslator

    MAX_CHUNK = 4500
    chunks = []
    remaining = text
    while len(remaining) > MAX_CHUNK:
        cut = remaining[:MAX_CHUNK].rfind('\n')
        if cut == -1:
            cut = remaining[:MAX_CHUNK].rfind('. ')
        if cut == -1:
            cut = MAX_CHUNK
        chunks.append(remaining[:cut + 1])
        remaining = remaining[cut + 1:]
    if remaining:
        chunks.append(remaining)

    translator = GoogleTranslator(source=source_lang, target="es")
    translated_chunks = []
    for c in chunks:
        if c.strip():
            translated_chunks.append(translator.translate(c))
    return "\n".join(translated_chunks)


# ─────────────────────────────────────────────
# TRADUCCION: IA (Gemini / Claude)
# ─────────────────────────────────────────────

def _gemini(text, is_google_result):
    try:
        from google import genai
    except ImportError:
        raise Exception("Ejecuta: pip install google-genai")

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise Exception(
            "Falta GEMINI_API_KEY en .env\n"
            "Obt\u00e9n una clave GRATIS en: https://aistudio.google.com/app/apikey"
        )

    client = genai.Client(api_key=api_key)

    rules = (
        "REGLAS OBLIGATORIAS DE FORMATO ESPA\u00d1OL:\n"
        "1. DI\u00c1LOGOS: Usa gui\u00f3n largo (\u2014) SIEMPRE, NUNCA comillas. "
        "Formato: \u2014Texto del di\u00e1logo \u2014dijo el personaje\u2014. Continuaci\u00f3n.\n"
        "2. INCISOS DEL NARRADOR: \u2014Hola \u2014dijo Juan\u2014. Luego se fue. "
        "(gui\u00f3n largo antes y despu\u00e9s del inciso, punto despu\u00e9s del cierre).\n"
        "3. SIGNOS DE APERTURA: Siempre usar \u00bf...? y \u00a1...! (con apertura).\n"
        "4. MAY\u00daSCULAS: Despu\u00e9s de . ? ! y ... siempre may\u00fascula.\n"
        "5. PUNTUACI\u00d3N: Sin espacio antes de . , ; : ? ! \u2014 Espacio despu\u00e9s.\n"
        "6. P\u00c1RRAFOS: Mant\u00e9n los saltos de p\u00e1rrafo del original. "
        "Cada di\u00e1logo de distinto personaje en p\u00e1rrafo aparte.\n"
        "7. NATURALIDAD: Usa expresiones naturales del espa\u00f1ol castellano, "
        "no traducciones literales. Adapta modismos e interjecciones.\n"
        "8. REGISTRO: Mant\u00e9n el registro del original (formal/informal, t\u00fa/usted).\n"
        "9. NOMBRES PROPIOS: No traduzcas nombres propios de personas ni lugares ficticios.\n"
        "10. COMAS VOCATIVAS: Siempre coma antes del vocativo: \u2014Hola, Juan.\n"
        "11. LE\u00cdSMO/LA\u00cdSMO: Usa los pronombres correctamente seg\u00fan la RAE.\n"
        "12. PUNTOS SUSPENSIVOS: Exactamente tres puntos (...), espacio despu\u00e9s.\n"
    )

    if is_google_result:
        prompt = (
            "Eres un corrector y editor literario experto en espa\u00f1ol castellano.\n"
            "Se te proporciona una traducci\u00f3n autom\u00e1tica de un libro/texto largo.\n"
            "Tu tarea es MEJORARLA para que suene como prosa literaria profesional "
            "publicada en Espa\u00f1a.\n\n"
            f"{rules}\n"
            "CORRECCIONES ADICIONALES A BUSCAR:\n"
            "- Frases que suenan a traducci\u00f3n literal del ingl\u00e9s\n"
            "- Gerundios mal usados (en espa\u00f1ol se usan menos que en ingl\u00e9s)\n"
            "- Voz pasiva excesiva (en espa\u00f1ol se prefiere activa o pasiva refleja)\n"
            "- Repeticiones innecesarias de sujeto pronominable\n"
            "- Falsos amigos (actually\u2260actualmente, etc.)\n\n"
            "Devuelve \u00daNICAMENTE el texto mejorado, sin explicaciones ni comentarios.\n\n"
            f"TEXTO A MEJORAR:\n{text}"
        )
    else:
        prompt = (
            "Eres un traductor literario profesional especializado en espa\u00f1ol castellano.\n"
            "Traduce el siguiente texto manteniendo el estilo, tono y ritmo narrativo del original.\n"
            "La traducci\u00f3n debe leerse como una obra publicada en espa\u00f1ol, no como una traducci\u00f3n.\n\n"
            f"{rules}\n"
            "INSTRUCCIONES ADICIONALES:\n"
            "- Adapta modismos e interjecciones al espa\u00f1ol natural\n"
            "- Evita gerundios excesivos, voz pasiva innecesaria y calcos del ingl\u00e9s\n"
            "- Mant\u00e9n la estructura de p\u00e1rrafos del original\n"
            "- Para cada di\u00e1logo de personaje distinto, usa p\u00e1rrafo aparte\n\n"
            "Devuelve \u00daNICAMENTE la traducci\u00f3n, sin explicaciones ni comentarios.\n\n"
            f"TEXTO ORIGINAL:\n{text}"
        )

    response = client.models.generate_content(
        model="gemini-2.0-flash",
        contents=prompt,
    )
    return response.text


def _claude(text, is_google_result):
    try:
        import anthropic
    except ImportError:
        raise Exception("Ejecuta: pip install anthropic")

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise Exception("Falta ANTHROPIC_API_KEY en .env")

    client = anthropic.Anthropic(api_key=api_key)

    rules = (
        "REGLAS: Di\u00e1logos con gui\u00f3n largo (\u2014), nunca comillas. "
        "Incisos: \u2014texto \u2014dijo X\u2014. Contin\u00faa. "
        "Signos de apertura obligatorios: \u00bf...? \u00a1...! "
        "May\u00fascula despu\u00e9s de . ? ! y puntos suspensivos. "
        "Coma vocativa. Sin gerundios ni pasiva innecesarios. "
        "Espa\u00f1ol castellano natural, no traducci\u00f3n literal."
    )

    if is_google_result:
        action = (
            f"Mejora esta traducci\u00f3n autom\u00e1tica para que suene como prosa literaria "
            f"profesional en espa\u00f1ol castellano. {rules}"
        )
    else:
        action = (
            f"Traduce al espa\u00f1ol castellano como prosa literaria publicada. "
            f"Mant\u00e9n estilo, tono y ritmo. {rules}"
        )
    prompt = (
        f"{action}\n"
        f"Devuelve \u00daNICAMENTE el texto resultante, sin explicaciones.\n\nTEXTO:\n{text}"
    )
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text


# ─────────────────────────────────────────────
# PDF <-> DOCX CONVERSION
# ─────────────────────────────────────────────

def _pdf_to_docx(pdf_path, docx_path):
    """Convert PDF to DOCX preserving layout, fonts, images."""
    from pdf2docx import Converter
    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()


def _docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX back to PDF via MS Word / LibreOffice."""
    from docx2pdf import convert
    convert(str(docx_path), str(pdf_path))


def _extract_paragraphs(docx_path):
    """Extract all non-empty paragraphs from DOCX."""
    from docx import Document
    doc = Document(str(docx_path))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
    return paragraphs


def _translate_docx_in_place(docx_path, output_path, source_lang="auto",
                              use_ai=False, provider="gemini",
                              progress_callback=None):
    """
    Translate a DOCX file in-place preserving all formatting.
    Each paragraph's text is translated while keeping runs/styles intact.
    Returns stats dict.
    """
    from docx import Document
    doc = Document(str(docx_path))

    total = sum(1 for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) >= 3)
    translated_count = 0

    for para in doc.paragraphs:
        original_text = para.text.strip()
        if not original_text or len(original_text) < 3:
            continue

        try:
            translated = _translate_text_google(original_text, source_lang)

            if use_ai:
                try:
                    if provider == "gemini":
                        translated = _gemini(translated, is_google_result=True)
                    else:
                        translated = _claude(translated, is_google_result=True)
                except Exception as ai_err:
                    print(f"AI refinement failed, using Google result: {ai_err}")

            translated = post_process_spanish(translated)
            _apply_translation_to_paragraph(para, translated)
            translated_count += 1

            if progress_callback:
                progress_callback(translated_count, total)

        except Exception as e:
            print(f"Error translating paragraph: {e}")

    doc.save(str(output_path))
    return {"total_paragraphs": total, "translated": translated_count}


def _apply_translation_to_paragraph(para, translated_text):
    """Replace text in a paragraph while preserving first-run formatting."""
    runs = para.runs
    if not runs:
        return
    runs[0].text = translated_text
    for run in runs[1:]:
        run.text = ""


# ─────────────────────────────────────────────
# NEW API: DOCX PIPELINE
# ─────────────────────────────────────────────

@app.route("/api/upload", methods=["POST"])
def upload_pdf():
    """Upload PDF and auto-convert to DOCX (preserves layout)."""
    if "file" not in request.files:
        return jsonify({"error": "No se envi\u00f3 ning\u00fan archivo"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "El archivo debe ser un PDF"}), 400

    session_id = str(uuid.uuid4())
    session_path = get_session_path(session_id)

    pdf_path = session_path / "original.pdf"
    file.save(str(pdf_path))

    # Page count
    import fitz
    doc = fitz.open(str(pdf_path))
    total_pages = len(doc)
    doc.close()

    # Auto-convert PDF -> DOCX
    docx_path = session_path / "original.docx"
    try:
        _pdf_to_docx(pdf_path, docx_path)
    except Exception as e:
        return jsonify({"error": f"Error convirtiendo PDF a DOCX: {str(e)}"}), 500

    paragraphs = _extract_paragraphs(docx_path)

    meta = {
        "session_id": session_id,
        "original_filename": file.filename,
        "total_pages": total_pages,
        "total_paragraphs": len(paragraphs),
        "status": "uploaded",
    }
    save_session_meta(session_id, meta)

    return jsonify({
        "session_id": session_id,
        "filename": file.filename,
        "total_pages": total_pages,
        "total_paragraphs": len(paragraphs),
        "paragraphs_preview": paragraphs[:20],
    })


@app.route("/api/paragraphs/<session_id>")
def get_paragraphs(session_id):
    """Return all paragraphs from the DOCX."""
    session_path = get_session_path(session_id)
    docx_path = session_path / "original.docx"
    if not docx_path.exists():
        return jsonify({"error": "DOCX no encontrado"}), 404
    paragraphs = _extract_paragraphs(docx_path)
    return jsonify({"paragraphs": paragraphs, "count": len(paragraphs)})


@app.route("/api/translate-docx", methods=["POST"])
def translate_docx():
    """
    Translate all paragraphs in the DOCX in-place.
    Preserves formatting, images, layout from original.
    """
    data = request.json
    session_id = data.get("session_id")
    source_lang = data.get("source_lang", "auto")
    use_ai = data.get("use_ai", False)
    provider = data.get("provider", "gemini")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404

    session_path = get_session_path(session_id)
    docx_path = session_path / "original.docx"
    translated_docx = session_path / "translated.docx"

    if not docx_path.exists():
        return jsonify({"error": "DOCX original no encontrado"}), 404

    try:
        stats = _translate_docx_in_place(
            docx_path, translated_docx,
            source_lang=source_lang,
            use_ai=use_ai,
            provider=provider,
        )

        meta["status"] = "translated"
        meta["translation_method"] = f"google{'+' + provider if use_ai else ''}"
        save_session_meta(session_id, meta)

        return jsonify({
            "success": True,
            "total_paragraphs": stats["total_paragraphs"],
            "translated": stats["translated"],
            "method": meta["translation_method"],
        })
    except Exception as e:
        return jsonify({"error": f"Error traduciendo: {str(e)}"}), 500


@app.route("/api/generate-pdf", methods=["POST"])
def generate_pdf():
    """Convert the translated DOCX back to PDF (preserves layout)."""
    data = request.json
    session_id = data.get("session_id")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404

    session_path = get_session_path(session_id)
    translated_docx = session_path / "translated.docx"

    if not translated_docx.exists():
        return jsonify({"error": "Primero debes traducir el documento"}), 400

    output_filename = f"traduccion_{session_id[:8]}.pdf"
    output_path = session_path / output_filename

    try:
        _docx_to_pdf(translated_docx, output_path)

        meta["status"] = "completed"
        meta["output_filename"] = output_filename
        save_session_meta(session_id, meta)

        return jsonify({
            "success": True,
            "filename": output_filename,
            "download_url": f"/api/download/{session_id}/{output_filename}",
        })
    except Exception as e:
        return jsonify({"error": f"Error generando PDF: {str(e)}"}), 500


# ─────────────────────────────────────────────
# LEGACY ENDPOINTS (backward compat)
# ─────────────────────────────────────────────

@app.route("/api/split", methods=["POST"])
def split_pdf():
    data = request.json
    session_id = data.get("session_id")
    pages_per_part = int(data.get("pages_per_part", 10))

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404

    session_path = get_session_path(session_id)
    pdf_path = session_path / "original.pdf"

    import fitz
    src = fitz.open(str(pdf_path))
    total_pages = len(src)

    parts_dir = session_path / "parts"
    if parts_dir.exists():
        shutil.rmtree(str(parts_dir))
    parts_dir.mkdir()

    parts = []
    part_num = 1
    for start in range(0, total_pages, pages_per_part):
        end = min(start + pages_per_part, total_pages)
        part_doc = fitz.open()
        part_doc.insert_pdf(src, from_page=start, to_page=end - 1)
        part_filename = f"parte_{part_num:03d}.pdf"
        part_doc.save(str(parts_dir / part_filename))
        part_doc.close()
        parts.append({
            "part_num": part_num,
            "filename": part_filename,
            "pages_start": start + 1,
            "pages_end": end,
            "page_count": end - start,
            "status": "pending",
            "translated_text": None,
        })
        part_num += 1
    src.close()

    meta["parts"] = parts
    meta["pages_per_part"] = pages_per_part
    meta["status"] = "split"
    save_session_meta(session_id, meta)
    return jsonify({"session_id": session_id, "total_parts": len(parts), "parts": parts})


@app.route("/api/extract-text", methods=["POST"])
def extract_text():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404

    session_path = get_session_path(session_id)
    part_path = session_path / "parts" / f"parte_{part_num:03d}.pdf"
    if not part_path.exists():
        return jsonify({"error": f"Parte {part_num} no encontrada"}), 404

    import fitz
    doc = fitz.open(str(part_path))
    text = ""
    for page in doc:
        text += page.get_text() + "\n\n"
    doc.close()

    return jsonify({"part_num": part_num, "text": text.strip(), "char_count": len(text)})


@app.route("/api/translate-google", methods=["POST"])
def translate_google():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    source_lang = data.get("source_lang", "auto")
    text = data.get("text", "")

    if not text.strip():
        return jsonify({"error": "No hay texto para traducir"}), 400

    try:
        translated = _translate_text_google(text, source_lang)
        translated = post_process_spanish(translated)

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = translated
                part["status"] = "translated_google"
                break
        save_session_meta(session_id, meta)

        return jsonify({"part_num": part_num, "translated_text": translated, "method": "google"})
    except Exception as e:
        return jsonify({"error": f"Error en traducci\u00f3n: {str(e)}"}), 500


@app.route("/api/translate-ai", methods=["POST"])
def translate_ai():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    text = data.get("text", "")
    is_google_result = data.get("is_google_result", False)
    provider = data.get("provider", "gemini")

    if not text.strip():
        return jsonify({"error": "No hay texto para procesar"}), 400

    try:
        if provider == "gemini":
            translated = _gemini(text, is_google_result)
        else:
            translated = _claude(text, is_google_result)
        translated = post_process_spanish(translated)

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = translated
                part["status"] = f"translated_{provider}"
                break
        save_session_meta(session_id, meta)

        return jsonify({"part_num": part_num, "translated_text": translated, "method": provider})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/translate-pipeline", methods=["POST"])
def translate_pipeline():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    source_lang = data.get("source_lang", "auto")
    text = data.get("text", "")
    skip_ai = data.get("skip_ai", False)
    provider = data.get("provider", "gemini")

    if not text.strip():
        return jsonify({"error": "No hay texto para traducir"}), 400

    steps_done = []
    try:
        google_result = _translate_text_google(text, source_lang)
        steps_done.append("google")

        ai_result = google_result
        if not skip_ai:
            try:
                if provider == "gemini":
                    ai_result = _gemini(google_result, is_google_result=True)
                else:
                    ai_result = _claude(google_result, is_google_result=True)
                steps_done.append(provider)
            except Exception as ai_err:
                print(f"AI failed, using Google: {ai_err}")

        final = post_process_spanish(ai_result)
        steps_done.append("postprocess")

        meta = load_session_meta(session_id)
        for part in meta.get("parts", []):
            if part["part_num"] == part_num:
                part["translated_text"] = final
                part["status"] = "translated_pipeline"
                break
        save_session_meta(session_id, meta)

        return jsonify({
            "part_num": part_num,
            "translated_text": final,
            "method": "+".join(steps_done),
            "steps": steps_done,
        })
    except Exception as e:
        return jsonify({"error": f"Error en pipeline: {str(e)}"}), 500


@app.route("/api/save-translation", methods=["POST"])
def save_translation():
    data = request.json
    session_id = data.get("session_id")
    part_num = int(data.get("part_num", 1))
    translated_text = data.get("translated_text", "")

    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404

    for part in meta.get("parts", []):
        if part["part_num"] == part_num:
            part["translated_text"] = translated_text
            part["status"] = "translated_manual"
            break

    save_session_meta(session_id, meta)
    return jsonify({"success": True, "part_num": part_num})


@app.route("/api/fix-dialogues", methods=["POST"])
def fix_dialogues_api():
    data = request.json
    fixed = post_process_spanish(data.get("text", ""))
    return jsonify({"fixed_text": fixed})


# ─────────────────────────────────────────────
# DOWNLOAD / SESSION
# ─────────────────────────────────────────────

@app.route("/api/download/<session_id>/<filename>")
def download_file(session_id, filename):
    fp = get_session_path(session_id) / filename
    if not fp.exists():
        return jsonify({"error": "Archivo no encontrado"}), 404
    return send_file(str(fp), as_attachment=True, download_name=filename)


@app.route("/api/session/<session_id>")
def get_session(session_id):
    meta = load_session_meta(session_id)
    if not meta:
        return jsonify({"error": "Sesi\u00f3n no encontrada"}), 404
    return jsonify(meta)


# ─────────────────────────────────────────────
# SERVIR FRONTEND
# ─────────────────────────────────────────────

FRONTEND_DIR = BASE_DIR / "frontend"

@app.route("/")
def serve_frontend():
    return send_file(str(FRONTEND_DIR / "index.html"))

@app.route("/<path:filename>")
def serve_static(filename):
    fp = FRONTEND_DIR / filename
    if fp.exists() and fp.is_file():
        return send_file(str(fp))
    return jsonify({"error": "No encontrado"}), 404


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    print(f"\n\U0001f680 PDF Translator v3 \u2014 http://localhost:{port}")
    print(f"   \u2705 Pipeline: PDF \u2192 DOCX \u2192 Traducir \u2192 PDF")
    print(f"   \u2705 Preserva layout, fuentes, im\u00e1genes, tablas")
    print(f"   \u2705 Google Translate + Gemini/Claude refinamiento")
    print(f"   \u2705 Post-procesamiento ling\u00fc\u00edstico espa\u00f1ol\n")
    app.run(debug=True, port=port)
